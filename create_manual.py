from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 設定預設字型 ──
style = doc.styles["Normal"]
font = style.font
font.name = "PingFang TC"
font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

# ── 顏色常數 ──
TEAL = RGBColor(0x0D, 0x73, 0x77)
DARK = RGBColor(0x1A, 0x2A, 0x2A)
GRAY = RGBColor(0x5A, 0x6F, 0x6F)

# ═══════════════════════════════════════
# 封面
# ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("【科技守護．智慧安居】")
run.font.size = Pt(16)
run.font.color.rgb = TEAL
run.font.bold = True
run.font.name = "PingFang TC"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2.add_run("長照人員科技輔具教育訓練\n報名系統使用手冊")
run.font.size = Pt(22)
run.font.color.rgb = DARK
run.font.bold = True
run.font.name = "PingFang TC"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("屏東縣輔具資源中心（屏中區／屏北區）\n115 年 4 月")
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.font.name = "PingFang TC"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

doc.add_page_break()

# ═══════════════════════════════════════
# 目錄
# ═══════════════════════════════════════
h = doc.add_heading("目錄", level=1)
for run in h.runs:
    run.font.color.rgb = TEAL

toc_items = [
    "一、系統概覽",
    "二、網址一覽",
    "三、前台：報名頁面操作說明",
    "四、後台：報名管理頁面操作說明",
    "五、DM 使用說明",
    "六、常見問題（FAQ）",
    "七、聯繫窗口",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════
# Helper
# ═══════════════════════════════════════
def add_heading_teal(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL
    return h

def add_url_box(label, url, note=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(f"{label}\n")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.name = "PingFang TC"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

    run = p.add_run(url)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = TEAL
    run.font.name = "PingFang TC"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

    if note:
        run = p.add_run(f"\n{note}")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.font.name = "PingFang TC"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

    # 加邊框底色
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): "F4F9F9",
        qn("w:val"): "clear",
    })
    shading.append(shd)
    doc.add_paragraph()

def add_step(num, title, desc):
    p = doc.add_paragraph()
    run = p.add_run(f"步驟 {num}：{title}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    run.font.name = "PingFang TC"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.left_indent = Cm(1)
    p2.paragraph_format.space_after = Pt(8)

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════
# 一、系統概覽
# ═══════════════════════════════════════
add_heading_teal("一、系統概覽")

doc.add_paragraph(
    "本系統為「科技守護．智慧安居」長照人員科技輔具教育訓練專用的線上報名平台，"
    "包含前台報名頁面與後台管理頁面兩大功能。系統部署於 Netlify 雲端平台，"
    "支援即時報名人數追蹤、剩餘名額顯示，以及報名名單的查詢與匯出。"
)

doc.add_paragraph()
add_heading_teal("系統架構", level=2)

table = doc.add_table(rows=4, cols=3)
table.style = "Light Grid Accent 1"

headers = ["功能", "說明", "對象"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h

rows_data = [
    ["報名頁面（前台）", "學員個人線上報名、選擇場次、查看剩餘名額", "報名學員"],
    ["管理後台", "即時名單查詢、場次篩選、匯出 CSV", "管理人員"],
    ["課程 DM", "活動宣傳單張，含報名連結", "對外發送"],
]
for i, row_data in enumerate(rows_data):
    for j, val in enumerate(row_data):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()

# ═══════════════════════════════════════
# 二、網址一覽
# ═══════════════════════════════════════
add_heading_teal("二、網址一覽")

doc.add_paragraph("本系統共有兩個頁面，網址如下：")
doc.add_paragraph()

add_url_box(
    "前台 — 報名頁面（提供給學員）",
    "https://ubiquitous-pasca-0a9bae.netlify.app/",
    "學員開啟此連結即可直接報名"
)

add_url_box(
    "後台 — 報名管理頁面（僅限管理人員）",
    "https://ubiquitous-pasca-0a9bae.netlify.app/admin.html",
    "需輸入管理密碼才能進入"
)

p = doc.add_paragraph()
run = p.add_run("重要提醒：")
run.font.bold = True
run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
p.add_run("後台網址請勿提供給學員，僅供屏中區、屏北區管理人員使用。")

# ═══════════════════════════════════════
# 三、前台操作說明
# ═══════════════════════════════════════
doc.add_page_break()
add_heading_teal("三、前台：報名頁面操作說明")

doc.add_paragraph("學員開啟報名網址後，依序完成以下欄位即可送出報名。")
doc.add_paragraph()

add_heading_teal("3.1 填寫報名資料", level=2)

table = doc.add_table(rows=7, cols=3)
table.style = "Light Grid Accent 1"

headers = ["欄位", "必填", "說明"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

fields = [
    ["姓名", "是", "填寫報名者本人姓名"],
    ["服務單位", "是", "下拉選單選擇單位；如不在列表中，選「其他」後自行輸入"],
    ["職稱", "是", "下拉選單選擇（A 個管）；如為其他職稱，選「其他」後自行輸入"],
    ["手機號碼", "是", "報名者手機"],
    ["Email", "否", "選填，供寄送確認通知"],
    ["選擇場次", "是", "4 個場次擇一；每場顯示即時剩餘名額，額滿則無法選擇"],
]
for i, row in enumerate(fields):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()

add_heading_teal("3.2 場次剩餘名額說明", level=2)

add_bullet("剩餘 11 位以上 → 綠色標籤「剩餘 XX 位」")
add_bullet("剩餘 6–10 位 → 橘色標籤「剩餘 XX 位」")
add_bullet("剩餘 1–5 位 → 橘色標籤「剩餘 XX 位（即將額滿）」")
add_bullet("剩餘 0 位 → 紅色標籤「已額滿」，該場次無法選擇")

doc.add_paragraph()

add_heading_teal("3.3 送出報名", level=2)

doc.add_paragraph(
    "確認資料無誤後，點擊「送出報名」按鈕。成功後會顯示報名成功頁面。"
    "每位學員僅需填寫一次，不需由單位統一彙報。"
)

# ═══════════════════════════════════════
# 四、後台操作說明
# ═══════════════════════════════════════
doc.add_page_break()
add_heading_teal("四、後台：報名管理頁面操作說明")

add_heading_teal("4.1 登入", level=2)

add_step(1, "開啟後台網址",
    "在瀏覽器輸入：https://ubiquitous-pasca-0a9bae.netlify.app/admin.html")
add_step(2, "輸入管理密碼",
    "預設密碼為 ptrc2026，輸入後按「登入」。")
add_step(3, "進入管理介面",
    "登入成功後會自動載入所有報名資料。")

doc.add_paragraph()

add_heading_teal("4.2 查看報名統計", level=2)

doc.add_paragraph(
    "頁面上方有 4 張場次統計卡片，每張顯示："
)
add_bullet("已報名人數 / 30（總名額）")
add_bullet("進度條（綠色 → 橘色 → 紅色，依報名比例變化）")

doc.add_paragraph()

add_heading_teal("4.3 篩選場次", level=2)

doc.add_paragraph(
    "點擊任一場次統計卡片，表格會自動篩選顯示該場次的報名名單。"
    "再點一次同張卡片可取消篩選，或點擊「清除篩選」按鈕回到全部。"
)

doc.add_paragraph()

add_heading_teal("4.4 匯出 CSV", level=2)

add_step(1, "選擇匯出範圍",
    "可先點擊場次卡片篩選特定場次，或不篩選匯出全部。")
add_step(2, "點擊「匯出 CSV」",
    "右上角橘色按鈕，點擊後自動下載 CSV 檔案。")
add_step(3, "開啟檔案",
    "下載的 CSV 可直接用 Excel 或 Numbers 開啟，包含完整報名資訊。")

doc.add_paragraph()

add_heading_teal("4.5 重新整理", level=2)

doc.add_paragraph(
    "點擊右上角「重新整理」按鈕，可即時取得最新報名資料。"
    "建議在活動報名期間定期重新整理以掌握最新報名狀況。"
)

# ═══════════════════════════════════════
# 五、DM 使用說明
# ═══════════════════════════════════════
doc.add_page_break()
add_heading_teal("五、DM 使用說明")

doc.add_paragraph(
    "專案中附有課程 DM 檔案（dm.html），可透過以下方式使用："
)

add_bullet("以瀏覽器開啟 dm.html，使用列印功能（Ctrl+P / Cmd+P）儲存為 PDF")
add_bullet("直接螢幕截圖後以圖片形式分享至 LINE 群組或 Email")
add_bullet("DM 底部已附上報名網址區域與 QR Code 預留位置")

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("提醒：")
run.font.bold = True
p.add_run("DM 中的報名網址已更新為正式網址。右下角 QR Code 區域目前為預留位置，"
          "可使用任何 QR Code 產生器（如 Google 搜尋「QR Code Generator」）"
          "將報名網址轉為 QR Code 圖片後替換。")

# ═══════════════════════════════════════
# 六、FAQ
# ═══════════════════════════════════════
doc.add_page_break()
add_heading_teal("六、常見問題（FAQ）")

faqs = [
    (
        "Q1：忘記後台密碼怎麼辦？",
        "預設密碼為 ptrc2026。如需更改密碼，請至 Netlify 後台 → "
        "Site configuration → Environment variables → 新增或修改 ADMIN_KEY 變數值。"
        "修改後需重新部署才會生效（可至 Deploys 頁面點擊 Trigger deploy）。"
    ),
    (
        "Q2：場次顯示「已額滿」但實際有人取消，如何處理？",
        "目前系統計數為累加機制，如有取消需手動調整。"
        "建議至 Netlify 後台 → Forms → training-registration 中刪除該筆提交，"
        "並聯繫系統管理員重置該場次計數。"
    ),
    (
        "Q3：可以用手機操作後台嗎？",
        "可以。後台頁面支援手機瀏覽器，但建議使用電腦操作體驗較佳，"
        "特別是在查看表格與匯出 CSV 時。"
    ),
    (
        "Q4：報名資料存在哪裡？安全嗎？",
        "報名資料儲存在 Netlify 雲端平台（Netlify Blobs + Netlify Forms），"
        "後台需密碼才能存取。資料不會公開，僅管理員可查看。"
    ),
    (
        "Q5：如何修改報名頁面上的單位名單？",
        "需編輯 index.html 中的 <select name=\"unit\"> 區塊，"
        "新增或移除 <option> 項目後推送至 GitHub，Netlify 會自動重新部署。"
    ),
    (
        "Q6：如何更改網站網址？",
        "至 Netlify 後台 → Site configuration → Site name，"
        "可將 ubiquitous-pasca-0a9bae 改為自訂名稱（如 ptrc-training），"
        "網址就會變成 ptrc-training.netlify.app。"
    ),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    run.font.name = "PingFang TC"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

    p2 = doc.add_paragraph(a)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(12)

# ═══════════════════════════════════════
# 七、聯繫窗口
# ═══════════════════════════════════════
add_heading_teal("七、聯繫窗口")

table = doc.add_table(rows=3, cols=3)
table.style = "Light Grid Accent 1"

headers = ["區域", "聯繫人", "電話"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

contacts = [
    ["屏中區", "陳社工", "(08) 789-9599 #22"],
    ["屏北區", "孫社工", "(08) 736-5455 #24"],
]
for i, row in enumerate(contacts):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("屏東縣輔具資源中心\n社團法人屏東縣輔具應用及身心健康促進協會")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.font.name = "PingFang TC"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

# ── 儲存 ──
output = "/Users/yangmingqin/Downloads/tech-training-registration/報名系統使用手冊.docx"
doc.save(output)
print(f"Saved: {output}")
